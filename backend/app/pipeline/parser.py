"""
Stage 2: Document Parser

Extracts text content from various document formats.
Each format has a dedicated parser that preserves layout structure.
"""

import io
import csv
from dataclasses import dataclass, field
from pathlib import Path

from app.config import settings


@dataclass
class ParseResult:
    """Result of document parsing."""
    text: str
    page_count: int = 1
    method: str = "direct"  # "direct", "ocr", "hybrid"
    warnings: list[str] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)


async def parse_pdf(file_path: str) -> ParseResult:
    """
    Parse PDF files with a hybrid approach:
    - Try direct text extraction first (fast, accurate for text PDFs)
    - Fall back to OCR for scanned/image PDFs (per-page detection)
    """
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    page_count = len(doc)
    all_text = []
    warnings = []
    method = "direct"
    ocr_pages = 0

    for page_num in range(page_count):
        page = doc[page_num]
        text = page.get_text("text")

        if text and text.strip() and len(text.strip()) > 20:
            # Page has extractable text
            all_text.append(f"--- Page {page_num + 1} ---\n{text.strip()}")
        else:
            # Page appears to be scanned — try OCR
            try:
                pix = page.get_pixmap(dpi=300)
                img_data = pix.tobytes("png")

                from PIL import Image
                import pytesseract

                img = Image.open(io.BytesIO(img_data))
                ocr_text = pytesseract.image_to_string(img)

                if ocr_text and ocr_text.strip():
                    all_text.append(f"--- Page {page_num + 1} (OCR) ---\n{ocr_text.strip()}")
                    ocr_pages += 1
                else:
                    warnings.append(f"Page {page_num + 1}: OCR returned no text (page may be blank or image too dark)")
            except Exception as e:
                warnings.append(f"Page {page_num + 1}: OCR failed — {str(e)}")

    doc.close()

    if ocr_pages > 0 and ocr_pages == page_count:
        method = "ocr"
    elif ocr_pages > 0:
        method = "hybrid"

    combined_text = "\n\n".join(all_text)

    return ParseResult(
        text=combined_text,
        page_count=page_count,
        method=method,
        warnings=warnings,
        metadata={"ocr_pages": ocr_pages},
    )


async def parse_image(file_path: str) -> ParseResult:
    """Parse image files using OCR with preprocessing."""
    from PIL import Image, ImageEnhance
    import pytesseract

    warnings = []

    try:
        img = Image.open(file_path)

        # Auto-rotate based on EXIF data
        try:
            from PIL import ExifTags
            exif = img.getexif()
            if exif:
                for tag_id, value in exif.items():
                    tag = ExifTags.TAGS.get(tag_id, tag_id)
                    if tag == "Orientation":
                        if value == 3:
                            img = img.rotate(180, expand=True)
                        elif value == 6:
                            img = img.rotate(270, expand=True)
                        elif value == 8:
                            img = img.rotate(90, expand=True)
        except Exception:
            pass  # EXIF rotation is best-effort

        # Enhance contrast for better OCR
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.5)

        # Convert to grayscale for OCR
        img = img.convert("L")

        text = pytesseract.image_to_string(img)

        if not text or not text.strip():
            warnings.append("OCR returned no text. Image may be blank, too dark, or not contain text.")

        return ParseResult(
            text=text.strip() if text else "",
            method="ocr",
            warnings=warnings,
            metadata={"image_size": f"{img.width}x{img.height}"},
        )
    except Exception as e:
        return ParseResult(
            text="",
            method="ocr",
            warnings=[f"Image parsing failed: {str(e)}"],
        )


async def parse_docx(file_path: str) -> ParseResult:
    """Parse DOCX files, extracting paragraphs and tables."""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            # Preserve heading structure
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                parts.append(f"{'#' * int(level)} {para.text.strip()}")
            else:
                parts.append(para.text.strip())

    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        parts.append(f"\n--- Table {table_idx + 1} ---")
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            parts.append(row_text)

    return ParseResult(
        text="\n\n".join(parts),
        method="direct",
        metadata={"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)},
    )


async def parse_csv(file_path: str) -> ParseResult:
    """Parse CSV files into a structured text representation."""
    import pandas as pd

    warnings = []
    try:
        # Try reading with pandas for robust CSV handling
        df = pd.read_csv(file_path, encoding="utf-8", on_bad_lines="warn")
    except UnicodeDecodeError:
        try:
            df = pd.read_csv(file_path, encoding="latin-1", on_bad_lines="warn")
            warnings.append("File was not UTF-8; decoded as Latin-1.")
        except Exception as e:
            return ParseResult(
                text="",
                warnings=[f"CSV parsing failed: {str(e)}"],
            )
    except Exception as e:
        # Try as TSV
        try:
            df = pd.read_csv(file_path, sep="\t", on_bad_lines="warn")
            warnings.append("Parsed as TSV (tab-separated).")
        except Exception:
            return ParseResult(
                text="",
                warnings=[f"CSV/TSV parsing failed: {str(e)}"],
            )

    # Convert to readable text
    header = " | ".join(str(c) for c in df.columns)
    rows = []
    for _, row in df.head(500).iterrows():  # Limit to 500 rows for processing
        rows.append(" | ".join(str(v) for v in row.values))

    text = f"Columns: {header}\n\n" + "\n".join(rows)

    if len(df) > 500:
        warnings.append(f"Only first 500 of {len(df)} rows included for processing.")

    return ParseResult(
        text=text,
        method="direct",
        warnings=warnings,
        metadata={"rows": len(df), "columns": len(df.columns)},
    )


async def parse_text(file_path: str) -> ParseResult:
    """Parse plain text and markdown files."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="latin-1") as f:
            text = f.read()

    return ParseResult(
        text=text,
        method="direct",
        metadata={"chars": len(text), "lines": text.count("\n") + 1},
    )


async def parse_xlsx(file_path: str) -> ParseResult:
    """Parse Excel files."""
    import pandas as pd

    warnings = []
    try:
        xlsx = pd.ExcelFile(file_path)
        parts = []

        for sheet_name in xlsx.sheet_names:
            df = pd.read_excel(xlsx, sheet_name=sheet_name)
            header = " | ".join(str(c) for c in df.columns)
            rows = [" | ".join(str(v) for v in row.values) for _, row in df.head(500).iterrows()]

            parts.append(f"--- Sheet: {sheet_name} ---\nColumns: {header}\n" + "\n".join(rows))

            if len(df) > 500:
                warnings.append(f"Sheet '{sheet_name}': only first 500 of {len(df)} rows included.")

        return ParseResult(
            text="\n\n".join(parts),
            method="direct",
            warnings=warnings,
            metadata={"sheets": len(xlsx.sheet_names)},
        )
    except Exception as e:
        return ParseResult(
            text="",
            warnings=[f"Excel parsing failed: {str(e)}"],
        )


# ── Parser router ──────────────────────────────────────────────

PARSER_MAP = {
    ".pdf": parse_pdf,
    ".png": parse_image,
    ".jpg": parse_image,
    ".jpeg": parse_image,
    ".tiff": parse_image,
    ".bmp": parse_image,
    ".docx": parse_docx,
    ".doc": parse_docx,  # Best-effort; .doc may not work with python-docx
    ".csv": parse_csv,
    ".xlsx": parse_xlsx,
    ".txt": parse_text,
    ".md": parse_text,
}


async def parse_document(file_path: str) -> ParseResult:
    """
    Route to the appropriate parser based on file extension.

    This is the main entry point for Stage 2 of the pipeline.
    """
    ext = Path(file_path).suffix.lower()
    parser = PARSER_MAP.get(ext)

    if not parser:
        return ParseResult(
            text="",
            warnings=[f"No parser available for file type: {ext}"],
        )

    return await parser(file_path)
